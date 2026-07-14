"""
Generador de documentos Word usando python-docx-template (Jinja2 en .docx).

Flujo:
  1. Carga la plantilla .docx base del tipo de informe (1-4)
  2. Renderiza con los datos del informe (contenido_json)
  3. Guarda el .docx en app/static/docx/
  4. Actualiza ruta_docx e incrementa version en la tabla informes
"""
from datetime import datetime, timezone
from pathlib import Path
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches
from loguru import logger
from sqlalchemy.orm import Session
from app.models.informe import Informe

PLANTILLAS_DIR = Path(__file__).parent.parent / "static" / "plantillas"
DOCX_OUTPUT_DIR = Path(__file__).parent.parent / "static" / "docx"
GRAFICOS_DIR = Path(__file__).parent.parent / "static" / "graficos"

PLANTILLAS_DIR.mkdir(parents=True, exist_ok=True)
DOCX_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
GRAFICOS_DIR.mkdir(parents=True, exist_ok=True)


def _ruta_plantilla(tipo_informe: int) -> Path:
    return PLANTILLAS_DIR / f"informe_{tipo_informe}_plantilla.docx"


def _asegurar_datos_caratula(db: Session, informe: Informe) -> dict:
    """
    Garantiza que el informe tenga los datos de la carátula, vengan de donde vengan.

    Un informe puede nacer vacío (p. ej. al guardar el checklist AVAC se crea la
    fila sin pasar por el generador), y entonces la carátula salía coja: "PERIODO"
    sin número y sin el texto de carreras. Aquí se rellena lo que falte leyéndolo
    de la base de datos.

    Solo rellena lo ausente: si el usuario editó un valor desde la UI, ese manda
    (la editabilidad total del informe es un requisito).
    """
    from sqlalchemy.orm.attributes import flag_modified
    from app.models.area import Area
    from app.models.consejo import ConsejoCarrera
    from app.models.jefatura import JefaturaArea
    from app.models.usuario import Usuario
    from app.services.generador_informes import _datos_caratula, nombre_para_firma

    contenido = dict(informe.contenido_json or {})

    consejo = db.query(ConsejoCarrera).filter(ConsejoCarrera.id == informe.consejo_id).first()
    periodo = consejo.periodo if consejo else None
    area = db.query(Area).filter(Area.id == informe.area_id).first()

    jefe = None
    if consejo:
        jefe = (
            db.query(Usuario)
            .join(JefaturaArea, JefaturaArea.usuario_id == Usuario.id)
            .filter(
                JefaturaArea.area_id == informe.area_id,
                JefaturaArea.periodo_id == consejo.periodo_id,
            )
            .first()
        )

    faltantes = {
        **_datos_caratula(periodo),
        "periodo_nombre": periodo.nombre if periodo else "",
        "area_nombre": area.nombre if area else "",
        "jefe_nombre": nombre_para_firma(jefe.nombre_completo) if jefe else "",
        "jefe_titulo": (jefe.titulo or "") if jefe else "",
    }

    cambio = False
    for clave, valor in faltantes.items():
        if not contenido.get(clave) and valor:
            contenido[clave] = valor
            cambio = True

    if cambio:
        informe.contenido_json = contenido
        flag_modified(informe, "contenido_json")   # sin esto la columna JSON no se guarda
        db.commit()
        logger.info(f"Informe {informe.id}: datos de carátula completados desde la BD")

    return contenido


def generar_docx(db: Session, informe: Informe) -> str:
    """
    Genera el .docx para el informe usando su contenido_json.
    Retorna la ruta relativa guardada en la DB.

    Si no existe plantilla real, genera un .docx básico de texto.
    """
    plantilla_ruta = _ruta_plantilla(informe.tipo_informe)
    contenido = _asegurar_datos_caratula(db, informe)

    nombre_archivo = (
        f"informe_{informe.tipo_informe}"
        f"_consejo{informe.consejo_id}"
        f"_area{informe.area_id}"
        f"_v{informe.version}.docx"
    )
    ruta_salida = DOCX_OUTPUT_DIR / nombre_archivo

    if plantilla_ruta.exists():
        graficos = _generar_desde_plantilla(plantilla_ruta, contenido, ruta_salida, informe)
        _guardar_rutas_graficos(db, informe, graficos)
    else:
        logger.warning(
            f"Plantilla {plantilla_ruta} no encontrada — generando .docx básico"
        )
        _generar_docx_basico(informe, contenido, ruta_salida)

    # Actualizar informe en DB
    informe.ruta_docx = nombre_archivo
    informe.generado_en = datetime.now(timezone.utc)
    db.commit()

    logger.info(f"Informe {informe.id} generado: {nombre_archivo}")
    return nombre_archivo


def _aplanar_contexto(contenido: dict) -> dict:
    """
    Prepara el contexto para Jinja2.

    Las plantillas usan marcadores planos (`{{ agenda }}`), pero el contenido se
    guarda anidado bajo `secciones` / `secciones_direccion`. Sin este aplanado,
    los informes salían con los títulos pero SIN texto.

    Las secciones de dirección se exponen con el prefijo `dir_` para poder
    imprimirlas junto al aporte del jefe de área.
    """
    contexto = dict(contenido)
    for campo, valor in (contenido.get("secciones") or {}).items():
        contexto.setdefault(campo, valor)
    for campo, valor in (contenido.get("secciones_direccion") or {}).items():
        contexto[f"dir_{campo}"] = valor
    return contexto


def _entero(valor) -> int:
    """Un conteo del contenido_json. Tras editarlo en la UI puede llegar como texto."""
    try:
        return int(float(valor))
    except (TypeError, ValueError):
        return 0


def _incrustar_graficos(doc: DocxTemplate, contexto: dict, informe: Informe) -> None:
    """
    Redibuja el gráfico de cada asignatura y lo incrusta. `{{ c.grafico }}`.

    El PNG se **vuelve a dibujar en cada generación**, a partir de los rangos que
    tenga el `contenido_json` en ese momento. Antes se guardaba una sola vez al
    crear el informe y el .docx reusaba ese archivo: si el usuario corregía el
    reparto alto/medio/bajo desde el editor, la tabla cambiaba pero el pastel
    seguía mostrando los números viejos.

    Si un item no da para gráfico (menos de dos rangos con datos), se deja el
    hueco vacío en lugar de romper la generación.
    """
    from app.services.graficos import grafico_rangos_interciclo, guardar_png

    originales = contexto.get("calificaciones_interciclo") or []
    if not originales:
        return

    # Copiar cada item: si mutáramos los originales, los objetos InlineImage
    # acabarían dentro de contenido_json y no son serializables a JSON.
    copias = []
    for idx, item in enumerate(originales):
        copia = dict(item)
        rangos = {
            "rango_alto": _entero(copia.get("rango_alto")),
            "rango_medio": _entero(copia.get("rango_medio")),
            "rango_bajo": _entero(copia.get("rango_bajo")),
        }
        titulo = f"{copia.get('asignatura', '')} — Grupo {copia.get('grupo', '')}".strip(" —")
        png = grafico_rangos_interciclo(rangos, titulo)

        if png:
            nombre = copia.get("grafico_ruta") or f"informe{informe.id}_interciclo_{idx}.png"
            guardar_png(png, GRAFICOS_DIR / nombre)
            copia["grafico_ruta"] = nombre
            copia["grafico"] = InlineImage(doc, str(GRAFICOS_DIR / nombre), width=Inches(5.4))
        else:
            copia["grafico"] = ""

        copias.append(copia)

    contexto["calificaciones_interciclo"] = copias


def _dos_decimales(valor):
    """
    Filtro Jinja `dec2`: fuerza dos decimales en las notas de las tablas.

    Se aplica al imprimir y no al calcular, porque los valores son editables: si
    el usuario corrige un promedio a "43.7", el informe igual debe mostrar 43.70.
    Lo que no sea número se devuelve tal cual (p. ej. un "—").
    """
    try:
        return f"{float(valor):.2f}"
    except (TypeError, ValueError):
        return valor


def _entorno_jinja() -> "Environment":
    from jinja2 import Environment

    env = Environment()
    env.filters["dec2"] = _dos_decimales
    return env


def _guardar_rutas_graficos(db: Session, informe: Informe, graficos: dict[int, str]) -> None:
    """
    Persiste en `contenido_json` el nombre del PNG de cada asignatura.

    Hace falta porque `GET /informes/{id}/grafico/{nombre}` solo sirve nombres
    declarados en el propio `contenido_json` (es lo que impide el path traversal).
    """
    if not graficos:
        return
    from sqlalchemy.orm.attributes import flag_modified

    contenido = dict(informe.contenido_json or {})
    items = [dict(x) for x in (contenido.get("calificaciones_interciclo") or [])]
    cambio = False
    for idx, nombre in graficos.items():
        if idx < len(items) and items[idx].get("grafico_ruta") != nombre:
            items[idx]["grafico_ruta"] = nombre
            cambio = True
    if cambio:
        contenido["calificaciones_interciclo"] = items
        informe.contenido_json = contenido
        flag_modified(informe, "contenido_json")


def _generar_desde_plantilla(
    plantilla_ruta: Path, contexto: dict, ruta_salida: Path, informe: Informe
) -> dict[int, str]:
    """
    Renderiza la plantilla Jinja2 .docx con el contexto dado.

    Devuelve {índice de asignatura: nombre del PNG} de los gráficos dibujados.
    """
    doc = DocxTemplate(str(plantilla_ruta))
    aplanado = _aplanar_contexto(contexto)
    _incrustar_graficos(doc, aplanado, informe)
    doc.render(aplanado, _entorno_jinja())
    doc.save(str(ruta_salida))

    return {
        idx: item["grafico_ruta"]
        for idx, item in enumerate(aplanado.get("calificaciones_interciclo") or [])
        if item.get("grafico_ruta")
    }


def _generar_docx_basico(informe: Informe, contenido: dict, ruta_salida: Path) -> None:
    """Genera un .docx básico cuando no hay plantilla disponible."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    TIPO_NOMBRES = {
        1: "Informe Inicial — Centro Docente",
        2: "Informe de Revisión AVAC",
        3: "Informe de Visitas Áulicas e Interciclo",
        4: "Informe Final de Calificaciones",
    }

    doc = Document()

    # Encabezado institucional
    encabezado = doc.add_heading(
        f"UNIVERSIDAD POLITÉCNICA SALESIANA — SEDE CUENCA", level=1
    )
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_heading("Carrera de Computación", level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nombre_tipo = TIPO_NOMBRES.get(informe.tipo_informe, f"Informe {informe.tipo_informe}")
    titulo = doc.add_heading(nombre_tipo.upper(), level=2)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Metadatos
    tabla_meta = doc.add_table(rows=3, cols=2)
    tabla_meta.style = "Table Grid"
    datos_meta = [
        ("Consejo de Carrera N°:", str(informe.consejo_id)),
        ("Área:", str(informe.area_id)),
        ("Versión:", str(informe.version)),
    ]
    for i, (etiqueta, valor) in enumerate(datos_meta):
        tabla_meta.rows[i].cells[0].text = etiqueta
        tabla_meta.rows[i].cells[1].text = valor

    doc.add_paragraph()

    # Contenido dinámico del informe
    secciones = contenido.get("secciones", {})
    if not secciones:
        doc.add_paragraph("Contenido del informe pendiente de generación.")
    else:
        for nombre_seccion, texto in secciones.items():
            doc.add_heading(nombre_seccion.replace("_", " ").title(), level=3)
            doc.add_paragraph(str(texto))

    # Análisis de calificaciones si existen
    analisis = contenido.get("analisis_calificaciones", [])
    if analisis:
        doc.add_heading("Análisis de Calificaciones", level=2)
        for item in analisis:
            doc.add_heading(
                f"Asignatura: {item.get('asignatura', '—')} — Grupo: {item.get('grupo', '—')}",
                level=3,
            )
            for campo, valor in item.items():
                if campo not in ("asignatura", "grupo"):
                    p = doc.add_paragraph()
                    p.add_run(f"{campo.replace('_', ' ').title()}: ").bold = True
                    p.add_run(str(valor))

    doc.add_paragraph()

    # Firmas
    doc.add_heading("Firmas", level=2)
    firmas = doc.add_table(rows=2, cols=2)
    firmas.rows[0].cells[0].text = "Jefe de Área"
    firmas.rows[0].cells[1].text = "Director/a de Carrera"
    firmas.rows[1].cells[0].text = "\n\n_______________________"
    firmas.rows[1].cells[1].text = "\n\n_______________________"

    doc.save(str(ruta_salida))


def regenerar_docx(db: Session, informe: Informe) -> str:
    """Incrementa la versión y regenera el .docx."""
    informe.version += 1
    db.flush()
    return generar_docx(db, informe)
