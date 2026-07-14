"""
Script para crear las plantillas .docx base con marcadores Jinja2.
Ejecutar UNA SOLA VEZ dentro del contenedor backend:

  docker compose exec backend python app/static/plantillas/crear_plantillas.py

Las plantillas se guardan en app/static/plantillas/ y son leídas
por doc_generator.py en tiempo de ejecución.
"""
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DIR = Path(__file__).parent
LOGO = DIR.parent / "logo-ups.jpg"  # app/static/logo-ups.jpg

# La carátula va en Cambria (con serifa, más institucional); todo el resto del
# informe, en Calibri. Es el formato que pidió el tutor.
CAMBRIA = "Cambria"
CALIBRI = "Calibri"

UPS_BLUE = RGBColor(0x00, 0x3D, 0xA5)
NEGRO = RGBColor(0x00, 0x00, 0x00)   # texto de la carátula
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
AZUL_HEX = "003DA5"        # relleno encabezados de tabla
GRIS_META_HEX = "E8EDF6"   # relleno columna de etiquetas (metadatos)


def _fijar_fuente(rpr, nombre: str):
    """
    Fija la tipografía en los cuatro scripts de Word (ascii/hAnsi/cs/eastAsia).

    Sin esto, `font.name` solo escribe `w:ascii` y Word rellena el resto con la
    fuente del tema, así que los encabezados salían en Cambria aunque el estilo
    dijera Calibri.
    """
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), nombre)


def _fuente_calibri(estilo, tamano: Pt, *, negrita: bool, color: RGBColor):
    """Aplica Calibri, sin cursiva, en un estilo (requisito del tutor)."""
    estilo.font.name = "Calibri"
    estilo.font.size = tamano
    estilo.font.bold = negrita
    estilo.font.italic = False
    estilo.font.color.rgb = color
    _fijar_fuente(estilo.element.get_or_add_rPr(), "Calibri")


def _aplicar_estilos(doc: Document):
    """Define estilos base uniformes para TODOS los informes.

    Así cada informe (1-4) se ve idéntico: misma tipografía (Calibri, sin
    cursiva), mismos tamaños y colores de encabezado, mismo interlineado.
    """
    normal = doc.styles["Normal"]
    _fuente_calibri(normal, Pt(11), negrita=False, color=GRIS_TEXTO)
    pf = normal.paragraph_format
    pf.space_after = Pt(10)   # 6pt dejaba los párrafos demasiado apretados
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Jerarquía tipográfica: 20 / 15 / 13 / 11 pt.
    # Antes el Heading 3 iba a 12pt y el cuerpo a 11pt — un solo punto de
    # diferencia, así que el subtítulo no se distinguía del texto.
    # El `space_after` separa el encabezado del párrafo que lleva debajo: con 2pt
    # el texto salía pegado al subtítulo.

    # Título del informe (Heading 1)
    h1 = doc.styles["Heading 1"]
    _fuente_calibri(h1, Pt(20), negrita=True, color=UPS_BLUE)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(16)

    # Secciones principales (Heading 2)
    h2 = doc.styles["Heading 2"]
    _fuente_calibri(h2, Pt(15), negrita=True, color=UPS_BLUE)
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(10)

    # Subsecciones (Heading 3)
    h3 = doc.styles["Heading 3"]
    _fuente_calibri(h3, Pt(13), negrita=True, color=RGBColor(0x22, 0x22, 0x22))
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(8)


def _no_estirar_saltos_de_linea(doc: Document):
    """
    Evita que Word estire las líneas que terminan en un salto manual.

    El cuerpo va justificado, y los valores con varias líneas (p. ej. las
    designaciones de jefes) llegan como un solo párrafo con saltos `<w:br/>`.
    Word justifica esas líneas igual que las demás y las abre hasta el margen,
    dejando huecos enormes entre palabras. `doNotExpandShiftReturn` las trata
    como última línea de párrafo, que es lo que son visualmente.
    """
    settings = doc.settings.element
    opcion = OxmlElement("w:doNotExpandShiftReturn")
    settings.append(opcion)


def _sombrear(cell, hexcolor: str):
    """Rellena el fondo de una celda con un color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _celda(cell, texto: str, *, negrita=False, blanco=False, relleno=None):
    """Escribe texto en una celda con formato uniforme, centrado verticalmente."""
    cell.text = texto
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    parrafo = cell.paragraphs[0]
    # El estilo Normal va justificado; dentro de una celda estrecha eso abre huecos
    # entre las palabras. En tablas se alinea a la izquierda.
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = parrafo.paragraph_format
    pf.line_spacing = 1.0          # 1.15 dejaba las filas demasiado altas
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)

    for run in parrafo.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = negrita
        run.font.italic = False
        run.font.color.rgb = BLANCO if blanco else GRIS_TEXTO
    if relleno:
        _sombrear(cell, relleno)


def _fijar_anchos(tabla, anchos):
    """
    Fija el ancho de cada columna.

    Word ignora los anchos si la tabla está en modo autoajuste, así que hay que
    desactivarlo y declarar `tblLayout=fixed`; sin eso, la columna del valor se
    estiraba hasta la mitad de la página aunque solo contuviera "48.10".
    El ancho hay que ponerlo celda por celda: en OOXML el ancho vive en la celda.
    """
    tabla.autofit = False
    tblPr = tabla._tbl.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    # Margen interno de las celdas: sin esto el texto queda pegado a los bordes.
    # (Unidades: twips — 1440 por pulgada.)
    margenes = OxmlElement("w:tblCellMar")
    for lado, valor in (("top", 60), ("bottom", 60), ("left", 110), ("right", 110)):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:w"), str(valor))
        el.set(qn("w:type"), "dxa")
        margenes.append(el)
    tblPr.append(margenes)

    for fila in tabla.rows:
        for i, ancho in enumerate(anchos):
            fila.cells[i].width = ancho


def _espacio(doc: Document, puntos: int = 10):
    """Un respiro en blanco (p. ej. entre un cuadro y el texto que le sigue)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.font.size = Pt(puntos)
    return p


def _meta(doc: Document, filas):
    """Tabla de metadatos uniforme: etiqueta (sombreada) | valor."""
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for etiqueta, marcador in filas:
        fila = tabla.add_row()
        _celda(fila.cells[0], etiqueta, negrita=True, relleno=GRIS_META_HEX)
        _celda(fila.cells[1], marcador)
    _fijar_anchos(tabla, (Inches(2.3), Inches(4.2)))
    doc.add_paragraph()
    return tabla


# Ancho útil de la página: 8.5in de papel − 2 × 1in de margen = 6.5in.
# La segunda columna solo lleva "Sí"/"No" o una nota con 2 decimales ("48.10"),
# así que se le da lo justo (1.1in) y el resto se lo queda la etiqueta.
ANCHOS_DATOS = (Inches(5.4), Inches(1.1))


def _tabla_datos(doc: Document, encabezados):
    """Crea una tabla con fila de encabezado azul + texto blanco."""
    tbl = doc.add_table(rows=1, cols=len(encabezados))
    tbl.style = "Table Grid"
    for i, h in enumerate(encabezados):
        _celda(tbl.rows[0].cells[i], h, negrita=True, blanco=True, relleno=AZUL_HEX)
    return tbl


def _firmas(doc: Document, columnas):
    """
    Bloque de firma: una línea sobre la que firmar, y debajo el nombre y el cargo.

    Sin tabla: una firma dentro de una celda con bordes se ve como un formulario,
    no como un documento firmado.
    """
    doc.add_paragraph()
    doc.add_heading("Firma de responsabilidad", level=2)

    for titulo, marcador in columnas:
        # Espacio en blanco para firmar a mano, y encima de la línea nada más
        hueco = doc.add_paragraph()
        hueco.paragraph_format.space_before = Pt(36)
        hueco.paragraph_format.space_after = Pt(0)

        linea = _parrafo(doc, "_______________________________________",
                         tamano=11, color=GRIS_TEXTO,
                         alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_despues=2)
        linea.paragraph_format.space_before = Pt(0)

        _parrafo(doc, marcador, tamano=11, negrita=True, color=GRIS_TEXTO,
                 espacio_despues=0)
        _parrafo(doc, titulo, tamano=10, color=GRIS_TEXTO, espacio_despues=0)


def _add_campo(parrafo, instr: str):
    """Inserta un campo de Word (PAGE / NUMPAGES) en un párrafo."""
    run = parrafo.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr_el)
    run._r.append(end)


def _configurar_seccion(doc: Document):
    """
    Header con logo UPS + footer con 'Carrera de Computación' y numeración.

    La carátula (página 1) NO lleva encabezado ni número de página: se usa
    `different_first_page_header_footer`, de modo que la numeración empieza en
    la página 2, como pide el formato institucional.

    Ni el encabezado ni el pie llevan línea divisoria: las únicas líneas del
    informe son las dos de la carátula.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # la carátula va limpia
    section.header_distance = Inches(0.5)

    # Encabezado y pie de la PRIMERA página: existen y van VACÍOS, a propósito.
    #
    # Con `titlePg` basta para que Word no pinte nada en la carátula, pero el
    # visor del navegador (docx-preview) no implementa esa marca: al no encontrar
    # un encabezado de primera página, caía al encabezado por defecto y pintaba el
    # logo ahí, encima del que ya lleva el cuerpo de la carátula (salía doble).
    # Declarándolos vacíos, el visor los usa y la vista previa coincide con Word.
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    # Encabezado (desde la página 2): logo a la IZQUIERDA, sin línea.
    #
    # El aire entre el logo y el primer título se hace con el `space_after` del
    # propio encabezado, NO subiendo el margen superior de la sección: el margen
    # aplica también a la carátula y empujaba su logo demasiado abajo.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(18)
    if LOGO.exists():
        hp.add_run().add_picture(str(LOGO), width=Inches(2.4))

    # Pie de página (desde la página 2): carrera (izq) y número de página (der)
    footer = section.footer
    fp = footer.paragraphs[0]
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Cm
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    fp.add_run("Carrera de Computación — Universidad Politécnica Salesiana, Sede Cuenca")
    fp.add_run("\tPág. ")
    _add_campo(fp, "PAGE")
    fp.add_run(" / ")
    _add_campo(fp, "NUMPAGES")


def _parrafo(doc, texto, *, tamano=11, negrita=False, color=GRIS_TEXTO,
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_antes=0, espacio_despues=6,
             fuente=CALIBRI):
    """Párrafo con formato directo, sin cursiva — usado en la carátula y el TOC."""
    p = doc.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(espacio_despues)
    run = p.add_run(texto)
    run.font.name = fuente
    run.font.size = Pt(tamano)
    run.font.bold = negrita
    run.font.italic = False
    run.font.color.rgb = color
    _fijar_fuente(run._r.get_or_add_rPr(), fuente)
    return p


# El ordinal del informe según su tipo (aparece en la carátula)
ORDINALES = {1: "PRIMER", 2: "SEGUNDO", 3: "TERCER", 4: "CUARTO"}


def _caratula(doc: Document, tipo: int):
    """
    Página 1 — carátula institucional, según el formato entregado por el tutor.

    El área, el período, las carreras y el jefe salen del sistema. No usa estilos
    de encabezado para que la carátula no aparezca en la tabla de contenidos.
    """
    ordinal = ORDINALES.get(tipo, "")

    # El logo va en el CUERPO de la carátula, no en el encabezado.
    # La página 1 lleva `titlePg` (sin encabezado, sin número de página), así que
    # el logo del encabezado —el que se ve de la página 2 en adelante— no aparece
    # aquí. Word lo omite correctamente; el visor del navegador (docx-preview) no
    # implementa `titlePg` y lo pintaba igual, y de ahí que la vista previa y el
    # Word descargado no coincidieran.
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.paragraph_format.space_after = Pt(36)
        p_logo.add_run().add_picture(str(LOGO), width=Inches(2.4))

    _parrafo(doc, f"{ordinal} INFORME DE", tamano=30, negrita=True, color=NEGRO,
             espacio_despues=0, fuente=CAMBRIA)
    _parrafo(doc, "JEFATURA DE ÁREA DE", tamano=30, negrita=True, color=NEGRO,
             espacio_despues=0, fuente=CAMBRIA)
    p = _parrafo(doc, "{{ area_nombre|upper }}", tamano=30, negrita=True, color=NEGRO,
                 espacio_despues=10, fuente=CAMBRIA)
    _borde_inferior(p)   # las únicas líneas del informe van aquí, en la carátula

    _parrafo(doc, "CORRESPONDIENTE AL PERIODO {{ periodo_numero }}",
             tamano=13, negrita=True, color=NEGRO, espacio_antes=8, espacio_despues=0,
             fuente=CAMBRIA)
    p = _parrafo(doc, "{{ carreras_texto }}", tamano=13, negrita=True, color=NEGRO,
                 espacio_despues=10, fuente=CAMBRIA)
    _borde_inferior(p)

    _parrafo(doc, "UNIVERSIDAD POLITÉCNICA SALESIANA",
             tamano=14, negrita=True, color=NEGRO, espacio_antes=40, fuente=CAMBRIA)
    _parrafo(doc, "SEDE CUENCA.", tamano=12, negrita=True, color=NEGRO,
             espacio_antes=10, fuente=CAMBRIA)

    _parrafo(doc, "{{ jefe_titulo }} {{ jefe_nombre }}",
             tamano=13, negrita=True, color=NEGRO, espacio_antes=90, espacio_despues=0,
             fuente=CAMBRIA)
    _parrafo(doc, "Jefe de Área de {{ area_nombre }}", tamano=9, color=NEGRO,
             fuente=CAMBRIA)

    doc.add_page_break()


def _tabla_contenidos(doc: Document):
    """
    Página 2 — tabla de contenidos.

    Se inserta un campo TOC real de Word. Word lo llena al abrir el documento
    (o con F9); por eso se deja un texto de respaldo dentro del campo.
    """
    _parrafo(doc, "TABLA DE CONTENIDOS", tamano=16, negrita=True, color=UPS_BLUE,
             alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=14)

    parrafo = doc.add_paragraph()
    run = parrafo.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")  # fuerza a Word a rellenarlo al abrir

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")

    texto = OxmlElement("w:t")
    texto.text = "Actualice esta tabla con clic derecho → Actualizar campos (F9)."

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, separador, texto, end):
        run._r.append(el)

    doc.add_page_break()


def _borde_inferior(parrafo):
    """Línea divisoria azul bajo el párrafo. Solo se usa en la carátula."""
    p = parrafo._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    borde = OxmlElement("w:bottom")
    borde.set(qn("w:val"), "single")
    borde.set(qn("w:sz"), "6")
    borde.set(qn("w:space"), "4")
    borde.set(qn("w:color"), "003DA5")
    pbdr.append(borde)
    p.append(pbdr)


def _encabezado_institucional(doc: Document, tipo: int, nombre: str):
    """
    Estructura común a los 4 informes:
      página 1 → carátula (sin numeración)
      página 2 → tabla de contenidos
      página 3 → título del informe y contenido
    """
    _aplicar_estilos(doc)         # Calibri sin cursiva, headings, interlineado
    _no_estirar_saltos_de_linea(doc)   # justificado sin huecos en los saltos manuales
    _configurar_seccion(doc)      # logo/pie desde la pág. 2; carátula limpia

    _caratula(doc, tipo)          # página 1
    _tabla_contenidos(doc)        # página 2

    n = doc.add_heading(f"INFORME {tipo} — {nombre.upper()}", level=1)
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def crear_informe_1():
    """Plantilla Informe 1 — Centro Docente (formulario dirección)."""
    doc = Document()
    _encabezado_institucional(doc, 1, "Centro Docente")

    # Metadatos
    _meta(doc, [
        ("Período académico:", "{{ periodo_nombre }}"),
        ("Área:", "{{ area_nombre }}"),
        ("Jefe de Área:", "{{ jefe_titulo }} {{ jefe_nombre }}"),
        ("Fecha del Consejo:", "{{ fecha_consejo }}"),
        ("Fecha del informe:", "{{ fecha_informe }}"),
    ])

    # Cada sección lleva: lo que escribió la dirección (común a todas las áreas)
    # y debajo el aporte propio del jefe de área.
    for titulo, campo in [
        ("1. Agenda tratada en la reunión", "agenda"),
        ("2. Designaciones de Jefes de Área", "designaciones"),
        ("3. Observaciones curriculares de docentes", "observaciones_curriculares"),
        ("4. Resultados de encuestas estudiantiles", "resultados_encuestas"),
        ("5. Resoluciones y compromisos", "resoluciones"),
        ("6. Observaciones adicionales", "observaciones_adicionales"),
    ]:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(f"{{{{ dir_{campo} }}}}")
        doc.add_paragraph(
            f"{{%p if {campo} %}}"
        )
        doc.add_heading("Aporte del Área", level=3)
        doc.add_paragraph(f"{{{{ {campo} }}}}")
        doc.add_paragraph("{%p endif %}")

    _firmas(doc, [
        ("Jefe de Área de {{ area_nombre }}", "{{ jefe_titulo }} {{ jefe_nombre }}"),
    ])

    ruta = DIR / "informe_1_plantilla.docx"
    doc.save(str(ruta))
    print(f"Creada: {ruta}")


def crear_informe_2():
    """Plantilla Informe 2 — Revisión AVAC."""
    doc = Document()
    _encabezado_institucional(doc, 2, "Revisión AVAC")

    _meta(doc, [
        ("Período:", "{{ periodo_nombre }}"),
        ("Área:", "{{ area_nombre }}"),
        ("Jefe de Área:", "{{ jefe_nombre }}"),
    ])

    doc.add_heading("Resultados del Checklist AVAC", level=2)
    doc.add_paragraph(
        "Los siguientes parámetros fueron evaluados para cada docente y asignatura:"
    )
    doc.add_paragraph("{% for item in checklists %}")
    doc.add_heading("{{ item.docente }} — {{ item.asignatura }} (Grupo {{ item.grupo }})", level=3)

    params = [
        ("Sílabo cargado", "silabo_cargado"),
        ("Registro de avance del sílabo", "registro_avance"),
        ("Guía de componente práctico", "guia_practicas"),
        ("Enlace consejería académica", "consejeria_academica"),
        ("Recursos con derechos de autor", "recursos_derechos_autor"),
        ("Libros digitales biblioteca", "libros_digitales"),
        ("Sección PRÁCTICAS", "seccion_practicas"),
        ("Guías de componente práctico", "guias_componente"),
        ("Actividades con rúbrica", "actividades_con_rubrica"),
        ("Sección INVESTIGATIVAS", "seccion_investigativas"),
        ("Actividad de investigación", "actividad_investigacion"),
        ("Proyecto integrador", "proyecto_integrador"),
    ]
    tbl = _tabla_datos(doc, ["Parámetro", "Cumple"])
    for nombre, campo in params:
        fila = tbl.add_row()
        _celda(fila.cells[0], nombre)
        _celda(fila.cells[1], f"{{{{ 'Sí' if item.{campo} else 'No' }}}}")
    _fijar_anchos(tbl, ANCHOS_DATOS)
    doc.add_paragraph()
    # Sin `if`: el generador nunca deja estos campos vacíos. Cuando no hay nada
    # escrito pone un texto explícito ("No se registraron observaciones."), para que
    # el informe no tenga huecos en blanco que se confundan con un olvido.
    doc.add_heading("Observaciones", level=3)
    doc.add_paragraph("{{ item.observaciones }}")
    doc.add_heading("Acciones de mejora del jefe de área", level=3)
    doc.add_paragraph("{{ item.acciones_mejora }}")
    # Generadas por IA a partir de los parámetros incumplidos y las observaciones
    doc.add_heading("Acciones de mejora sugeridas", level=3)
    doc.add_paragraph("{{ item.acciones_sugeridas }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_heading("Análisis de cumplimiento del área", level=2)
    doc.add_paragraph("{{ analisis_area }}")
    doc.add_heading("Acciones generales sugeridas para el área", level=3)
    doc.add_paragraph("{{ acciones_generales_avac }}")

    _firmas(doc, [
        ("Jefe de Área de {{ area_nombre }}", "{{ jefe_titulo }} {{ jefe_nombre }}"),
    ])

    ruta = DIR / "informe_2_plantilla.docx"
    doc.save(str(ruta))
    print(f"Creada: {ruta}")


def crear_informe_3():
    """Plantilla Informe 3 — Visitas Áulicas + Calificaciones Interciclo."""
    doc = Document()
    _encabezado_institucional(doc, 3, "Visitas Áulicas e Interciclo")

    _meta(doc, [
        ("Período:", "{{ periodo_nombre }}"),
        ("Área:", "{{ area_nombre }}"),
        ("Jefe de Área:", "{{ jefe_nombre }}"),
    ])

    doc.add_heading("PARTE A — Visitas Áulicas", level=2)
    doc.add_paragraph("{% for v in visitas %}")
    doc.add_heading("{{ v.docente }} — {{ v.asignatura }} (Grupo {{ v.grupo }})", level=3)
    tbl = _tabla_datos(doc, ["Parámetro", "Cumple"])
    for nombre, campo in [
        ("Visita realizada", "visita_realizada"),
        ("Puntualidad del docente", "puntualidad_docente"),
        ("Cumplimiento del sílabo", "cumplimiento_silabo"),
        ("Cumplimiento de prácticas", "cumplimiento_practicas"),
        ("Actividades con rúbrica", "actividades_con_rubrica"),
        ("Actividad de investigación", "actividad_investigacion"),
    ]:
        fila = tbl.add_row()
        _celda(fila.cells[0], nombre)
        _celda(fila.cells[1], f"{{{{ 'Sí' if v.{campo} else 'No' }}}}")
    _fijar_anchos(tbl, ANCHOS_DATOS)
    _espacio(doc)   # el texto no puede arrancar pegado al cuadro
    doc.add_paragraph("Observaciones estudiantes: {{ v.observaciones_estudiantes }}")
    doc.add_paragraph("Observaciones docente: {{ v.observaciones_docente }}")
    doc.add_paragraph(
        "Acciones de mejora del jefe de área al docente: {{ v.acciones_docente }}"
    )
    doc.add_paragraph("{% endfor %}")

    doc.add_heading("PARTE B — Calificaciones Interciclo (Parcial 1)", level=2)
    doc.add_paragraph("{% for c in calificaciones_interciclo %}")
    doc.add_heading("{{ c.asignatura }} — Grupo {{ c.grupo }}", level=3)
    tbl2 = _tabla_datos(doc, ["Indicador", "Valor"])
    # Las notas se imprimen con el filtro `dec2` (siempre 2 decimales: 48.10, 43.00).
    # Los conteos de estudiantes van enteros: "25.00 estudiantes" no significa nada.
    for etiqueta, campo, decimales in [
        ("Total estudiantes", "total_estudiantes", False),
        ("Nota máxima /50", "maximo", True),
        ("Nota mínima /50", "minimo", True),
        ("Promedio /50", "promedio", True),
        ("Mediana /50", "mediana", True),
        ("Rango alto (≥40)", "rango_alto", False),
        ("Rango medio (30-39)", "rango_medio", False),
        ("Rango bajo (<30)", "rango_bajo", False),
    ]:
        fila = tbl2.add_row()
        _celda(fila.cells[0], etiqueta, negrita=True, relleno=GRIS_META_HEX)
        filtro = "|dec2" if decimales else ""
        _celda(fila.cells[1], f"{{{{ c.{campo}{filtro} }}}}")
    _fijar_anchos(tbl2, ANCHOS_DATOS)
    _espacio(doc)

    # Gráfico de pastel con la distribución por rango de desempeño
    doc.add_paragraph("{%p if c.grafico %}")
    p_graf = doc.add_paragraph()
    p_graf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_graf.paragraph_format.space_before = Pt(6)
    p_graf.paragraph_format.space_after = Pt(14)   # aire bajo el gráfico
    p_graf.add_run("{{ c.grafico }}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("Análisis narrativo: {{ c.analisis_narrativo }}")
    doc.add_paragraph("Acciones de mejora: {{ c.acciones_mejora }}")
    # Aportes registrados por el docente sobre la materia (sumativos).
    # Sin `if`: si no registró nada, el generador pone un texto explícito.
    doc.add_heading("Observaciones del Docente a la Materia", level=3)
    doc.add_paragraph("{{ c.observaciones_materia }}")
    doc.add_heading("Acciones de Mejora propuestas por el Docente", level=3)
    doc.add_paragraph("{{ c.acciones_mejora_docente }}")
    doc.add_paragraph("{% endfor %}")

    _firmas(doc, [
        ("Jefe de Área de {{ area_nombre }}", "{{ jefe_titulo }} {{ jefe_nombre }}"),
    ])

    ruta = DIR / "informe_3_plantilla.docx"
    doc.save(str(ruta))
    print(f"Creada: {ruta}")


def crear_informe_4():
    """Plantilla Informe 4 — Análisis Final de Calificaciones."""
    doc = Document()
    _encabezado_institucional(doc, 4, "Análisis Final de Calificaciones")

    _meta(doc, [
        ("Período:", "{{ periodo_nombre }}"),
        ("Área:", "{{ area_nombre }}"),
        ("Jefe de Área:", "{{ jefe_nombre }}"),
    ])

    doc.add_paragraph("{% for c in calificaciones_finales %}")
    doc.add_heading("{{ c.docente }} — {{ c.asignatura }} (Grupo {{ c.grupo }})", level=2)

    sub_analisis = [
        ("1. Análisis general", "analisis_general"),
        ("2. Distribución aprobación/reprobación", "distribucion_aprobacion"),
        ("3. Comportamiento notas finales", "comportamiento_notas_finales"),
        ("4. Análisis Parcial 1", "analisis_parcial1"),
        ("5. Análisis Parcial 2", "analisis_parcial2"),
        ("6. Comparación entre parciales", "comparacion_parciales"),
        ("7. Uso de recuperación", "uso_recuperacion"),
        ("8. Relación parciales-nota final", "relacion_parciales_nota_final"),
        ("9. Identificación de outliers", "outliers"),
        ("10. Patrones generales", "patrones_generales"),
    ]
    for titulo, campo in sub_analisis:
        doc.add_heading(titulo, level=3)
        doc.add_paragraph(f"{{{{ c.{campo} }}}}")

    doc.add_heading("Acciones de mejora sugeridas", level=3)
    doc.add_paragraph("{{ c.acciones_mejora }}")

    # Aportes registrados por el docente sobre la materia (sumativos).
    # Sin `if`: si no registró nada, el generador pone un texto explícito.
    doc.add_heading("Observaciones del Docente a la Materia", level=3)
    doc.add_paragraph("{{ c.observaciones_materia }}")
    doc.add_heading("Acciones de Mejora propuestas por el Docente", level=3)
    doc.add_paragraph("{{ c.acciones_mejora_docente }}")
    doc.add_paragraph()
    doc.add_paragraph("{% endfor %}")

    doc.add_heading("Análisis Consolidado del Área", level=2)
    doc.add_paragraph("{{ analisis_consolidado_area }}")
    doc.add_heading("Acciones generales del área", level=3)
    doc.add_paragraph("{{ acciones_generales_area }}")

    _firmas(doc, [
        ("Jefe de Área de {{ area_nombre }}", "{{ jefe_titulo }} {{ jefe_nombre }}"),
    ])

    ruta = DIR / "informe_4_plantilla.docx"
    doc.save(str(ruta))
    print(f"Creada: {ruta}")


if __name__ == "__main__":
    crear_informe_1()
    crear_informe_2()
    crear_informe_3()
    crear_informe_4()
    print("Todas las plantillas creadas exitosamente.")
